import argparse
import hashlib
import json
import sqlite3
import sys
from pathlib import Path
from typing import Iterator

import requests
from docx import Document
from pypdf import PdfReader


OLLAMA_EMBED_URL = "http://127.0.0.1:11434/api/embed"
EMBEDDING_MODEL = "nomic-embed-text"

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOCUMENTS_ROOT = PROJECT_ROOT / "data" / "documents"
VECTORSTORE_ROOT = PROJECT_ROOT / "data" / "vectorstore"
DATABASE_PATH = VECTORSTORE_ROOT / "knowledge.db"

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
}

CHUNK_SIZE = 1_200
CHUNK_OVERLAP = 200


def create_database() -> sqlite3.Connection:
    VECTORSTORE_ROOT.mkdir(parents=True, exist_ok=True)

    connection = sqlite3.connect(DATABASE_PATH)

    connection.execute(
        """
        CREATE TABLE IF NOT EXISTS document_chunks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            file_path TEXT NOT NULL,
            file_name TEXT NOT NULL,
            department TEXT,
            page_number INTEGER,
            chunk_number INTEGER NOT NULL,
            content TEXT NOT NULL,
            content_hash TEXT NOT NULL,
            embedding TEXT NOT NULL,
            indexed_at TEXT DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(file_path, content_hash)
        )
        """
    )

    connection.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_document_chunks_file_path
        ON document_chunks(file_path)
        """
    )

    connection.commit()
    return connection


def clean_text(text: str) -> str:
    lines = []

    for line in text.replace("\x00", "").splitlines():
        cleaned_line = " ".join(line.split())

        if cleaned_line:
            lines.append(cleaned_line)

    return "\n".join(lines).strip()


def read_pdf(file_path: Path) -> Iterator[tuple[int | None, str]]:
    reader = PdfReader(str(file_path))

    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        cleaned = clean_text(text)

        if cleaned:
            yield page_index, cleaned


def read_docx(file_path: Path) -> Iterator[tuple[int | None, str]]:
    document = Document(str(file_path))
    sections: list[str] = []

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)

        if text:
            sections.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                clean_text(cell.text)
                for cell in row.cells
                if clean_text(cell.text)
            )

            if row_text:
                sections.append(row_text)

    combined_text = "\n".join(sections).strip()

    if combined_text:
        yield None, combined_text


def read_text_file(file_path: Path) -> Iterator[tuple[int | None, str]]:
    try:
        text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = file_path.read_text(encoding="cp1252", errors="replace")

    cleaned = clean_text(text)

    if cleaned:
        yield None, cleaned


def extract_sections(file_path: Path) -> Iterator[tuple[int | None, str]]:
    extension = file_path.suffix.lower()

    if extension == ".pdf":
        yield from read_pdf(file_path)
    elif extension == ".docx":
        yield from read_docx(file_path)
    elif extension in {".txt", ".md"}:
        yield from read_text_file(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}")


def split_into_chunks(
    text: str,
    chunk_size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP,
) -> list[str]:
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end]

        if end < len(text):
            preferred_break = max(
                chunk.rfind("\n"),
                chunk.rfind(". "),
                chunk.rfind("; "),
            )

            if preferred_break > chunk_size // 2:
                end = start + preferred_break + 1
                chunk = text[start:end]

        chunk = chunk.strip()

        if chunk:
            chunks.append(chunk)

        if end >= len(text):
            break

        start = max(end - overlap, start + 1)

    return chunks


def create_embedding(text: str) -> list[float]:
    try:
        response = requests.post(
            OLLAMA_EMBED_URL,
            json={
                "model": EMBEDDING_MODEL,
                "input": text,
            },
            timeout=300,
        )

        response.raise_for_status()
    except requests.ConnectionError as exc:
        raise RuntimeError(
            "Ollama is not running at 127.0.0.1:11434."
        ) from exc
    except requests.RequestException as exc:
        raise RuntimeError(
            f"Ollama embedding request failed: {exc}"
        ) from exc

    data = response.json()
    embeddings = data.get("embeddings")

    if not embeddings or not embeddings[0]:
        raise RuntimeError("Ollama returned no embedding.")

    return embeddings[0]


def get_department(file_path: Path) -> str:
    try:
        relative_path = file_path.relative_to(DOCUMENTS_ROOT)
        return relative_path.parts[0] if relative_path.parts else ""
    except ValueError:
        return ""


def hash_content(content: str) -> str:
    return hashlib.sha256(content.encode("utf-8")).hexdigest()


def find_supported_files(folder: Path) -> list[Path]:
    return sorted(
        file_path
        for file_path in folder.rglob("*")
        if file_path.is_file()
        and file_path.suffix.lower() in SUPPORTED_EXTENSIONS
        and not file_path.name.startswith("~$")
    )


def remove_existing_file_chunks(
    connection: sqlite3.Connection,
    relative_file_path: str,
) -> None:
    connection.execute(
        "DELETE FROM document_chunks WHERE file_path = ?",
        (relative_file_path,),
    )


def index_folder(folder: Path) -> None:
    if not folder.exists():
        raise FileNotFoundError(f"Folder not found: {folder}")

    files = find_supported_files(folder)

    if not files:
        print("No supported documents were found.")
        print("Supported types: .pdf, .docx, .txt, .md")
        return

    connection = create_database()

    total_files = len(files)
    indexed_files = 0
    indexed_chunks = 0
    skipped_files = 0

    print(f"\nFolder: {folder}")
    print(f"Database: {DATABASE_PATH}")
    print(f"Embedding model: {EMBEDDING_MODEL}")
    print(f"Supported files found: {total_files}\n")

    for file_number, file_path in enumerate(files, start=1):
        try:
            relative_path = str(
                file_path.relative_to(DOCUMENTS_ROOT)
            )

            print(
                f"[{file_number}/{total_files}] "
                f"Reading: {relative_path}"
            )

            extracted_sections = list(extract_sections(file_path))

            if not extracted_sections:
                print("    Skipped: no readable text was extracted.")
                skipped_files += 1
                continue

            remove_existing_file_chunks(connection, relative_path)

            file_chunk_number = 0

            for page_number, section_text in extracted_sections:
                chunks = split_into_chunks(section_text)

                for chunk in chunks:
                    file_chunk_number += 1

                    embedding = create_embedding(chunk)
                    content_hash = hash_content(chunk)

                    connection.execute(
                        """
                        INSERT OR REPLACE INTO document_chunks (
                            file_path,
                            file_name,
                            department,
                            page_number,
                            chunk_number,
                            content,
                            content_hash,
                            embedding
                        )
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            relative_path,
                            file_path.name,
                            get_department(file_path),
                            page_number,
                            file_chunk_number,
                            chunk,
                            content_hash,
                            json.dumps(embedding),
                        ),
                    )

                    indexed_chunks += 1
                    print(
                        f"    Indexed chunk {file_chunk_number}",
                        end="\r",
                        flush=True,
                    )

            connection.commit()
            indexed_files += 1

            print(
                f"    Completed: {file_chunk_number} chunks"
                + (" " * 20)
            )

        except Exception as exc:
            connection.rollback()
            skipped_files += 1
            print(f"    ERROR: {exc}")

    connection.close()

    print("\nIndexing complete.")
    print(f"Files indexed: {indexed_files}")
    print(f"Files skipped: {skipped_files}")
    print(f"Chunks created: {indexed_chunks}")
    print(f"Local database: {DATABASE_PATH}")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Index local company documents using Ollama."
    )

    parser.add_argument(
        "--folder",
        required=True,
        help=(
            "Folder name underneath data\\documents, "
            'such as "Accounts Receivable".'
        ),
    )

    args = parser.parse_args()

    target_folder = DOCUMENTS_ROOT / args.folder

    try:
        index_folder(target_folder)
    except Exception as exc:
        print(f"\nIndexing failed: {exc}")
        sys.exit(1)


if __name__ == "__main__":
    main()